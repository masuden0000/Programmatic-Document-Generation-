# Sistem Otomasi Template Dokumen - Streamlit Version
# Tampilan sederhana mirip ui.html

import hashlib
import json
import os
import pickle
import re
import sys
import tempfile
import time
from io import BytesIO
from typing import Any, Dict, List

import streamlit as st

# Configure page FIRST before any other streamlit calls
st.set_page_config(
    page_title="Sistem Otomasi Template Dokumen",
    page_icon="📄",
    layout="centered",
    initial_sidebar_state="collapsed",
)

from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH as WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Inches, Pt, RGBColor
from dotenv import load_dotenv
from langchain.schema import HumanMessage, SystemMessage
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import Docx2txtLoader
from langchain_google_genai import ChatGoogleGenerativeAI

# Load environment variables
load_dotenv()


class FormatRulesExtractor:
    """Mengekstrak aturan format dari dokumen panduan"""

    def __init__(self):
        # Initialize Gemini model
        # Try Streamlit secrets first, fallback to .env
        try:
            api_key = st.secrets["GOOGLE_API_KEY"]
        except:
            api_key = os.getenv("GOOGLE_API_KEY")

        if not api_key:
            st.error(
                "⚠️ GOOGLE_API_KEY tidak ditemukan! Pastikan sudah setting di Streamlit Cloud secrets atau file .env"
            )
            st.info("📝 Untuk Streamlit Cloud: Tambahkan GOOGLE_API_KEY di app secrets")
            st.stop()

        self.llm = ChatGoogleGenerativeAI(
            model="gemini-2.0-flash-exp",
            google_api_key=api_key,
            temperature=0.1,
            max_retries=3,
            timeout=60,
        )

        # Cache directory untuk menyimpan hasil
        self.cache_dir = "cache"
        os.makedirs(self.cache_dir, exist_ok=True)

    def _get_cache_key(self, text: str) -> str:
        """Generate cache key dari text content"""
        return hashlib.md5(text.encode()).hexdigest()

    def _load_from_cache(self, cache_key: str) -> Dict[str, Any] | None:
        """Load hasil dari cache jika ada"""
        cache_file = os.path.join(self.cache_dir, f"{cache_key}.pkl")
        if os.path.exists(cache_file):
            try:
                with open(cache_file, "rb") as f:
                    return pickle.load(f)
            except:
                pass
        return None

    def _save_to_cache(self, cache_key: str, data: Dict[str, Any]):
        """Simpan hasil ke cache"""
        cache_file = os.path.join(self.cache_dir, f"{cache_key}.pkl")
        try:
            with open(cache_file, "wb") as f:
                pickle.dump(data, f)
        except:
            pass

    def extract_rules_with_ai(self, document_path: str) -> Dict[str, Any]:
        """Menggunakan AI untuk ekstrak aturan format"""

        # Load dokumen
        try:
            if document_path.lower().endswith(".txt"):
                # Handle .txt files
                with open(document_path, "r", encoding="utf-8") as f:
                    full_text = f.read()
                st.markdown(
                    '<p style="color: #000000; background-color: #d1ecf1; padding: 10px; border-radius: 5px; border: 1px solid #bee5eb;">📄 File .txt berhasil dimuat</p>',
                    unsafe_allow_html=True,
                )
            else:
                # Handle .docx files
                loader = Docx2txtLoader(document_path)
                documents = loader.load()
                full_text = "\\n".join([doc.page_content for doc in documents])
                st.markdown(
                    '<p style="color: #000000; background-color: #d1ecf1; padding: 10px; border-radius: 5px; border: 1px solid #bee5eb;">📄 File .docx berhasil dimuat</p>',
                    unsafe_allow_html=True,
                )
        except Exception as e:
            st.error(f"❌ Error loading document: {str(e)}")
            return self._create_fallback_rules()

        # Cek cache terlebih dahulu
        cache_key = self._get_cache_key(full_text)
        cached_result = self._load_from_cache(cache_key)
        if cached_result:
            st.markdown(
                '<p style="color: #000000; background-color: #d1ecf1; padding: 10px; border-radius: 5px; border: 1px solid #bee5eb;">📋 Menggunakan hasil dari cache...</p>',
                unsafe_allow_html=True,
            )
            return cached_result

        # Jika teks terlalu panjang, potong ke 12000 karakter pertama
        if len(full_text) > 12000:
            full_text = full_text[:12000] + "..."
            st.warning(f"📏 Dokumen dipotong ke {len(full_text)} karakter pertama")

        st.markdown(
            f'<p style="color: #000000; background-color: #d1ecf1; padding: 10px; border-radius: 5px; border: 1px solid #bee5eb;">📊 Memproses {len(full_text)} karakter teks...</p>',
            unsafe_allow_html=True,
        )

        # Prompt yang lebih komprehensif dan detail
        prompt = f"""
        Analisis dokumen panduan berikut dan ekstrak SEMUA aturan format yang ada untuk membuat template dokumen yang LENGKAP dan DETAIL.

        DOKUMEN PANDUAN:
        {full_text}

        Ekstrak informasi berikut dalam format JSON yang SANGAT DETAIL:

        1. MARGIN (dalam cm): top, bottom, left, right
        2. FONT: family, size untuk berbagai elemen (body, heading, subheading)
        3. SPACING: line_spacing, paragraph_spacing, before_after_spacing
        4. PAPER: size (A4/Letter), orientation (portrait/landscape)
        5. HEADERS_FOOTERS: apakah ada, posisi, format
        6. NUMBERING: page_numbering, chapter_numbering, section_numbering
        7. DOCUMENT_STRUCTURE: urutan bagian dokumen (cover, toc, chapters, etc)
        8. SPECIAL_FORMATTING: indent, alignment, bullet_points, tables

        PENTING: Berikan nilai DEFAULT jika tidak disebutkan:
        - Margin: top=3, bottom=3, left=3, right=3
        - Font: family="Times New Roman", size=12
        - Spacing: line_spacing=1.5, paragraph_spacing=6
        - Paper: size="A4", orientation="portrait"

        Berikan response dalam format JSON yang valid dan lengkap.
        """

        try:
            # Buat messages untuk Gemini
            messages = [
                SystemMessage(
                    content="Anda adalah ahli analisis dokumen yang sangat detail dan teliti dalam mengekstrak aturan format."
                ),
                HumanMessage(content=prompt),
            ]

            # Kirim request ke Gemini dengan retry
            response = None
            for attempt in range(3):
                try:
                    st.markdown(
                        f'<p style="color: #000000; background-color: #d1ecf1; padding: 10px; border-radius: 5px; border: 1px solid #bee5eb;">🤖 Mengirim ke Gemini AI... (Percobaan {attempt + 1}/3)</p>',
                        unsafe_allow_html=True,
                    )
                    response = self.llm.invoke(messages)
                    st.markdown(
                        '<p style="color: #000000; background-color: #d4edda; padding: 10px; border-radius: 5px; border: 1px solid #c3e6cb;">✅ Response diterima dari Gemini AI</p>',
                        unsafe_allow_html=True,
                    )
                    break
                except Exception as e:
                    if attempt == 2:
                        st.error(f"❌ Semua percobaan gagal: {str(e)}")
                        raise e
                    st.warning(f"⚠️ Percobaan {attempt + 1} gagal, mencoba lagi...")
                    time.sleep(2)  # Wait before retry

            if not response:
                raise Exception("No response from Gemini API")

            # Parse response
            response_text = response.content

            # Ensure response_text is string
            if not isinstance(response_text, str):
                response_text = str(response_text)

            # Coba extract JSON dari response
            json_start = response_text.find("{")
            json_end = response_text.rfind("}") + 1

            if json_start != -1 and json_end != -1:
                json_text = response_text[json_start:json_end]
                try:
                    rules = json.loads(json_text)
                    st.markdown(
                        '<p style="color: #000000; background-color: #d4edda; padding: 10px; border-radius: 5px; border: 1px solid #c3e6cb;">✅ JSON berhasil di-parse</p>',
                        unsafe_allow_html=True,
                    )
                except json.JSONDecodeError as e:
                    st.warning(f"⚠️ JSON parsing gagal: {str(e)}")
                    rules = self._create_fallback_rules()
            else:
                st.warning(
                    "⚠️ JSON tidak ditemukan dalam response, menggunakan fallback rules"
                )
                rules = self._create_fallback_rules()

            # Normalisasi dan validasi rules
            rules = self._normalize_rules(rules)

            # Simpan ke cache
            self._save_to_cache(cache_key, rules)
            st.markdown(
                '<p style="color: #000000; background-color: #d4edda; padding: 10px; border-radius: 5px; border: 1px solid #c3e6cb;">💾 Hasil disimpan ke cache</p>',
                unsafe_allow_html=True,
            )

            return rules

        except Exception as e:
            st.warning(f"⚠️ AI extraction error: {str(e)}")
            # Return fallback rules
            return self._create_fallback_rules()

    def _create_fallback_rules(self) -> Dict[str, Any]:
        """Buat aturan format default jika AI gagal"""
        return {
            "margin": {"top": 3, "bottom": 3, "left": 3, "right": 3},
            "font": {
                "family": "Times New Roman",
                "body_size": 12,
                "heading_size": 14,
                "subheading_size": 12,
            },
            "spacing": {"line_spacing": 1.5, "paragraph_spacing": 6},
            "paper": {"size": "A4", "orientation": "portrait"},
            "headers_footers": {"enabled": True, "page_numbers": True},
            "numbering": {
                "page_numbering": "arabic",
                "chapter_numbering": "roman_upper",
                "section_numbering": "decimal",
            },
            "document_structure": [
                "Halaman Judul",
                "Daftar Isi",
                "BAB I PENDAHULUAN",
                "BAB II TINJAUAN PUSTAKA",
                "BAB III METODOLOGI",
                "BAB IV HASIL DAN PEMBAHASAN",
                "BAB V KESIMPULAN",
                "Daftar Pustaka",
            ],
        }

    def _normalize_rules(self, rules: Dict[str, Any]) -> Dict[str, Any]:
        """Normalisasi dan validasi rules"""
        normalized = self._create_fallback_rules()

        # Update dengan rules yang diekstrak
        if isinstance(rules, dict):
            for key, value in rules.items():
                if key in normalized:
                    if isinstance(value, dict) and isinstance(normalized[key], dict):
                        normalized[key].update(value)
                    else:
                        normalized[key] = value

        return normalized


class TemplateGenerator:
    """Generate template dokumen berdasarkan aturan format"""

    def __init__(self):
        """Initialize template generator with color utilities"""
        pass

    def _hex_to_rgb(self, hex_color: str) -> tuple:
        """Convert hex color to RGB tuple"""
        try:
            hex_color = hex_color.lstrip("#")
            if len(hex_color) == 6:
                return tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))
            return (0, 0, 0)  # Default to black
        except:
            return (0, 0, 0)  # Default to black

    def _is_light_color(self, rgb_color: tuple) -> bool:
        """Determine if a color is light (using relative luminance)"""
        r, g, b = rgb_color
        # Calculate relative luminance using sRGB formula
        luminance = (0.299 * r + 0.587 * g + 0.114 * b) / 255
        return luminance > 0.5

    def _get_font_color_for_background(self, bg_color: str) -> tuple:
        """Get appropriate font color (black or white) based on background"""
        if not bg_color or bg_color.lower() in ["white", "#ffffff", "#fff"]:
            return (0, 0, 0)  # Black for white/light backgrounds

        rgb = self._hex_to_rgb(bg_color)
        if self._is_light_color(rgb):
            return (0, 0, 0)  # Black for light backgrounds
        else:
            return (255, 255, 255)  # White for dark backgrounds

    def _apply_font_formatting(self, run, font_info: dict, bg_color: str = "#ffffff"):
        """Apply comprehensive font formatting to a run"""
        # Font family
        if "family" in font_info:
            run.font.name = font_info["family"]

        # Font size
        if "body_size" in font_info:
            run.font.size = Pt(font_info["body_size"])
        elif "size" in font_info:
            run.font.size = Pt(font_info["size"])

        # Font color based on background
        font_color = self._get_font_color_for_background(bg_color)
        run.font.color.rgb = RGBColor(*font_color)

        # Bold
        if font_info.get("bold", False):
            run.font.bold = True

        # Italic
        if font_info.get("italic", False):
            run.font.italic = True

        # Underline
        if font_info.get("underline", False):
            run.font.underline = True

    def _apply_paragraph_formatting(self, paragraph, spacing_info: dict):
        """Apply paragraph spacing and alignment"""
        paragraph_format = paragraph.paragraph_format

        # Line spacing
        line_spacing = spacing_info.get("line_spacing", 1.5)
        if isinstance(line_spacing, str):
            if "setengah" in line_spacing.lower() or "1.5" in line_spacing:
                line_spacing = 1.5
            elif "double" in line_spacing.lower() or "2" in line_spacing:
                line_spacing = 2.0
            else:
                line_spacing = 1.5
        paragraph_format.line_spacing = line_spacing

        # Space after paragraph
        space_after = spacing_info.get("after_paragraph", 6)
        if isinstance(space_after, str):
            try:
                space_after = float(re.findall(r"[\d.]+", space_after)[0])
            except:
                space_after = 6
        paragraph_format.space_after = Pt(space_after)

        # Space before paragraph
        space_before = spacing_info.get("before_paragraph", 0)
        if isinstance(space_before, str):
            try:
                space_before = float(re.findall(r"[\d.]+", space_before)[0])
            except:
                space_before = 0
        paragraph_format.space_before = Pt(space_before)

    def _apply_page_formatting(self, doc, rules: dict):
        """Apply page-level formatting (margins, orientation, page size, etc.)"""
        sections = doc.sections
        for section in sections:
            # Set page size to A4 explicitly
            section.page_width = Cm(21.0)  # A4 width: 21 cm
            section.page_height = Cm(29.7)  # A4 height: 29.7 cm

            # Page orientation
            orientation = rules.get("page_orientation", "portrait")
            if "landscape" in orientation.lower():
                section.orientation = WD_ORIENT.LANDSCAPE
                # Swap dimensions for landscape
                section.page_width = Cm(29.7)
                section.page_height = Cm(21.0)
            else:
                section.orientation = WD_ORIENT.PORTRAIT

            # Margins - Apply with more explicit conversion
            margins = rules.get("margin", {})

            # Top margin
            if "top" in margins:
                top_margin = margins["top"]
                if isinstance(top_margin, str):
                    try:
                        top_margin = float(re.findall(r"[\d.]+", top_margin)[0])
                    except:
                        top_margin = 3.0  # Default fallback
                section.top_margin = Cm(float(top_margin))

            # Bottom margin
            if "bottom" in margins:
                bottom_margin = margins["bottom"]
                if isinstance(bottom_margin, str):
                    try:
                        bottom_margin = float(re.findall(r"[\d.]+", bottom_margin)[0])
                    except:
                        bottom_margin = 3.0  # Default fallback
                section.bottom_margin = Cm(float(bottom_margin))

            # Left margin
            if "left" in margins:
                left_margin = margins["left"]
                if isinstance(left_margin, str):
                    try:
                        left_margin = float(re.findall(r"[\d.]+", left_margin)[0])
                    except:
                        left_margin = 3.0  # Default fallback
                section.left_margin = Cm(float(left_margin))

            # Right margin
            if "right" in margins:
                right_margin = margins["right"]
                if isinstance(right_margin, str):
                    try:
                        right_margin = float(re.findall(r"[\d.]+", right_margin)[0])
                    except:
                        right_margin = 3.0  # Default fallback
                section.right_margin = Cm(float(right_margin))

    def generate_template(self, rules: Dict[str, Any]) -> BytesIO:
        """Generate template Word document with comprehensive formatting"""
        try:
            # Create new document
            doc = Document()

            # Apply page formatting first
            self._apply_page_formatting(doc, rules)

            # Get formatting info
            font_info = rules.get("font", {})
            spacing_info = rules.get("spacing", {})
            bg_color = rules.get("background_color", "#ffffff")

            # Add title with special formatting
            title_text = rules.get("title", "TEMPLATE DOKUMEN")
            title = doc.add_heading(title_text, level=0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Format title
            for run in title.runs:
                title_font_info = font_info.copy()
                title_font_info["size"] = font_info.get(
                    "title_size", font_info.get("body_size", 12) + 2
                )
                title_font_info["bold"] = True
                self._apply_font_formatting(run, title_font_info, bg_color)

            # Apply spacing to title
            self._apply_paragraph_formatting(title, spacing_info)

            # Add document structure based on rules
            structure = rules.get("document_structure", [])

            for item in structure:
                if "BAB" in item.upper() or "CHAPTER" in item.upper():
                    # Main chapter heading
                    heading = doc.add_heading(item, level=1)
                    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    # Format chapter heading
                    for run in heading.runs:
                        chapter_font_info = font_info.copy()
                        chapter_font_info["size"] = font_info.get(
                            "heading_size", font_info.get("body_size", 12) + 1
                        )
                        chapter_font_info["bold"] = True
                        self._apply_font_formatting(run, chapter_font_info, bg_color)

                    self._apply_paragraph_formatting(heading, spacing_info)

                    # Add sample content paragraph
                    content_para = doc.add_paragraph(
                        "\\n[Konten untuk bagian ini akan diisi sesuai dengan panduan format yang telah diekstrak]\\n"
                    )

                    # Format content paragraph
                    for run in content_para.runs:
                        self._apply_font_formatting(run, font_info, bg_color)
                    self._apply_paragraph_formatting(content_para, spacing_info)

                else:
                    # Sub-heading
                    sub_heading = doc.add_heading(item, level=2)

                    # Format sub-heading
                    for run in sub_heading.runs:
                        sub_font_info = font_info.copy()
                        sub_font_info["size"] = font_info.get(
                            "subheading_size", font_info.get("body_size", 12)
                        )
                        sub_font_info["bold"] = True
                        self._apply_font_formatting(run, sub_font_info, bg_color)

                    self._apply_paragraph_formatting(sub_heading, spacing_info)

                    # Add sample content
                    content_para = doc.add_paragraph(
                        "\\n[Konten akan diisi di sini sesuai format yang diekstrak]\\n"
                    )

                    # Format content paragraph
                    for run in content_para.runs:
                        self._apply_font_formatting(run, font_info, bg_color)
                    self._apply_paragraph_formatting(content_para, spacing_info)

            # If no structure provided, add default sections
            if not structure:
                default_sections = [
                    "PENDAHULUAN",
                    "TINJAUAN PUSTAKA",
                    "METODOLOGI",
                    "HASIL DAN PEMBAHASAN",
                    "KESIMPULAN",
                ]

                for section in default_sections:
                    heading = doc.add_heading(section, level=1)
                    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    # Format heading
                    for run in heading.runs:
                        heading_font_info = font_info.copy()
                        heading_font_info["size"] = font_info.get(
                            "heading_size", font_info.get("body_size", 12) + 1
                        )
                        heading_font_info["bold"] = True
                        self._apply_font_formatting(run, heading_font_info, bg_color)

                    self._apply_paragraph_formatting(heading, spacing_info)

                    # Add content paragraph
                    content_para = doc.add_paragraph(
                        "\\n[Isi konten sesuai dengan aturan format yang diekstrak]\\n"
                    )

                    # Format content
                    for run in content_para.runs:
                        self._apply_font_formatting(run, font_info, bg_color)
                    self._apply_paragraph_formatting(content_para, spacing_info)

            # Add a formatting summary paragraph at the end
            doc.add_page_break()
            summary_heading = doc.add_heading(
                "RINGKASAN FORMAT YANG DITERAPKAN", level=1
            )
            summary_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Format summary heading
            for run in summary_heading.runs:
                summary_font_info = font_info.copy()
                summary_font_info["size"] = font_info.get(
                    "heading_size", font_info.get("body_size", 12) + 1
                )
                summary_font_info["bold"] = True
                self._apply_font_formatting(run, summary_font_info, bg_color)

            # Add formatting details
            format_details = []
            if font_info:
                format_details.append(f"Font: {font_info.get('family', 'Default')}")
                format_details.append(
                    f"Ukuran Font: {font_info.get('body_size', 12)} pt"
                )

            if rules.get("margin"):
                margins = rules["margin"]
                format_details.append(
                    f"Margin - Atas: {margins.get('top', 'default')} cm, Bawah: {margins.get('bottom', 'default')} cm"
                )
                format_details.append(
                    f"Margin - Kiri: {margins.get('left', 'default')} cm, Kanan: {margins.get('right', 'default')} cm"
                )

            if spacing_info:
                format_details.append(
                    f"Spasi Antar Baris: {spacing_info.get('line_spacing', 1.5)}"
                )
                format_details.append(
                    f"Spasi Setelah Paragraf: {spacing_info.get('after_paragraph', 6)} pt"
                )

            summary_text = (
                "\\n".join(format_details)
                if format_details
                else "Format default diterapkan."
            )
            summary_para = doc.add_paragraph(summary_text)

            # Format summary paragraph
            for run in summary_para.runs:
                self._apply_font_formatting(run, font_info, bg_color)
            self._apply_paragraph_formatting(summary_para, spacing_info)

            # Save to BytesIO
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            return doc_io

        except Exception as e:
            st.error(f"Error generating template: {str(e)}")
            # Return minimal template
            doc = Document()
            doc.add_heading("Template Dokumen", level=1)
            doc.add_paragraph("Template minimal - error dalam generating.")

            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            return doc_io


def main():
    # Simple CSS - mirip ui.html
    st.markdown(
        """
    <style>
    .stApp {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    }
    .block-container {
        background: white;
        border-radius: 20px;
        box-shadow: 0 20px 40px rgba(0, 0, 0, 0.1);
        max-width: 800px;
        padding: 2rem;
        margin-top: 2rem;
        margin-bottom: 2rem;
    }
    .header {
        text-align: center;
        margin-bottom: 2rem;
    }
    .header h1 {
        color: #667eea;
        font-size: 2.5em;
        margin-bottom: 0.5rem;
    }
    .header p {
        color: #666;
        font-size: 1.1em;
    }
    .upload-section {
        border: 3px dashed #667eea;
        border-radius: 15px;
        padding: 2rem;
        text-align: center;
        background: #f8f9ff;
        margin: 2rem 0;
    }
    .file-info {
        background: #f8f9ff;
        border: 1px solid #e0e0ff;
        border-radius: 10px;
        padding: 1rem;
        margin: 1rem 0;
        color: #000000;
    }
    .file-info strong {
        color: #000000;
    }
    .stButton > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        border-radius: 10px;
        padding: 0.75rem 2rem;
        font-weight: 600;
        width: 100%;
    }
    
    /* Custom styling for black text elements */
    .stSubheader {
        color: #000000 !important;
    }
    .stSubheader h3 {
        color: #000000 !important;
    }
    h3[id="upload-dokumen-panduan"] {
        color: #000000 !important;
    }
    
    /* Hide uploaded file display area completely */
    .st-emotion-cache-fis6aj {
        display: none !important;
    }
    .stFileUploaderFile {
        display: none !important;
    }
    div[data-testid="stFileUploaderFile"] {
        display: none !important;
    }
    .st-emotion-cache-14m29r0 {
        display: none !important;
    }
    .e16xj5sw6 {
        display: none !important;
    }
    .e16xj5sw5 {
        display: none !important;
    }
    
    /* Success and info alert messages - black text */
    .stAlert {
        color: #000000 !important;
    }
    .stAlert > div {
        color: #000000 !important;
    }
    .stAlert p {
        color: #000000 !important;
    }
    div[data-testid="stAlertContentSuccess"] {
        color: #000000 !important;
    }
    div[data-testid="stAlertContentSuccess"] p {
        color: #000000 !important;
    }
    div[data-testid="stAlertContentInfo"] {
        color: #000000 !important;
    }
    div[data-testid="stAlertContentInfo"] p {
        color: #000000 !important;
    }
    .st-emotion-cache-1w7qfeb p {
        color: #000000 !important;
    }
    
    /* Expander text styling */
    .stExpander {
        color: #000000 !important;
    }
    .stExpander > div {
        color: #000000 !important;
    }
    .stExpander summary {
        color: #000000 !important;
    }
    .stExpander details summary {
        color: #000000 !important;
    }
    .st-emotion-cache-1dtefog {
        color: #000000 !important;
    }
    .st-emotion-cache-1dtefog p {
        color: #000000 !important;
    }
    .st-emotion-cache-1ort0lt p {
        color: #000000 !important;
    }
    
    /* Info and success message containers */
    .stInfo {
        color: #000000 !important;
    }
    .stInfo > div {
        color: #000000 !important;
    }
    .stSuccess {
        color: #000000 !important;
    }
    .stSuccess > div {
        color: #000000 !important;
    }
    </style>
    """,
        unsafe_allow_html=True,
    )

    # Header - sederhana seperti ui.html
    st.markdown(
        """
    <div class="header">
        <h1>📄 Sistem Otomasi Template Dokumen</h1>
        <p>Ekstrak aturan format dari dokumen panduan dan buat template otomatis</p>
    </div>
    """,
        unsafe_allow_html=True,
    )

    # File upload
    st.markdown(
        '<h3 style="color: #000000;">📤 Upload Dokumen Panduan</h3>',
        unsafe_allow_html=True,
    )
    uploaded_file = st.file_uploader(
        "Pilih file dokumen panduan",
        type=["txt", "docx"],
        help="Upload dokumen yang berisi aturan format untuk template",
    )

    st.markdown("</div>", unsafe_allow_html=True)

    if uploaded_file is not None:
        # File info
        st.markdown(
            f"""
            <div class="file-info" style="color: #000000;">
                <strong style="color: #000000;">📁 File:</strong> {uploaded_file.name}<br>
                <strong style="color: #000000;">📏 Ukuran:</strong> {uploaded_file.size / 1024:.2f} KB
            </div>
            """,
            unsafe_allow_html=True,
        )

        # Save uploaded file temporarily
        try:
            temp_dir = tempfile.mkdtemp()
            temp_file_path = os.path.join(temp_dir, uploaded_file.name)

            with open(temp_file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())

            st.markdown(
                '<p style="color: #000000; background-color: #d4edda; padding: 10px; border-radius: 5px; border: 1px solid #c3e6cb;">✅ File berhasil disimpan: '
                + uploaded_file.name
                + "</p>",
                unsafe_allow_html=True,
            )
        except Exception as e:
            st.error(f"❌ Error menyimpan file: {str(e)}")
            st.stop()

        # Initialize components
        try:
            extractor = FormatRulesExtractor()
            generator = TemplateGenerator()
        except Exception as e:
            st.error(f"❌ Error initializing components: {str(e)}")
            st.stop()

        # Button columns
        col1, col2 = st.columns(2)

        with col1:
            if st.button("🔍 Ekstrak Aturan Format"):
                with st.spinner("Menganalisis dokumen..."):
                    try:
                        rules = extractor.extract_rules_with_ai(temp_file_path)
                        st.session_state.extracted_rules = rules
                        st.markdown(
                            '<p style="color: #000000; background-color: #d4edda; padding: 10px; border-radius: 5px; border: 1px solid #c3e6cb;">✅ Aturan format berhasil diekstrak!</p>',
                            unsafe_allow_html=True,
                        )

                        # Display rules
                        with st.expander(
                            "📋 Lihat Aturan yang Diekstrak", expanded=False
                        ):
                            st.markdown(
                                '<div style="color: #000000;">', unsafe_allow_html=True
                            )
                            st.json(rules)
                            st.markdown("</div>", unsafe_allow_html=True)

                    except Exception as e:
                        st.error(f"❌ Error: {str(e)}")

        with col2:
            if (
                st.button("📄 Generate Template")
                and "extracted_rules" in st.session_state
            ):
                with st.spinner("Membuat template..."):
                    try:
                        rules = st.session_state.extracted_rules
                        template_io = generator.generate_template(rules)

                        st.markdown(
                            '<p style="color: #000000; background-color: #d4edda; padding: 10px; border-radius: 5px; border: 1px solid #c3e6cb;">✅ Template berhasil dibuat!</p>',
                            unsafe_allow_html=True,
                        )

                        # Download button
                        st.download_button(
                            label="⬇️ Download Template",
                            data=template_io.getvalue(),
                            file_name="template_dokumen.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )

                    except Exception as e:
                        st.error(f"❌ Error: {str(e)}")

        # Cleanup
        try:
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
            if os.path.exists(temp_dir):
                os.rmdir(temp_dir)
        except Exception as e:
            # Silent cleanup - tidak perlu error jika gagal cleanup
            pass

    else:
        # Instructions when no file uploaded
        st.markdown(
            '<p style="color: #000000; background-color: #d1ecf1; padding: 10px; border-radius: 5px; border: 1px solid #bee5eb;">👆 Silakan upload dokumen panduan format untuk memulai</p>',
            unsafe_allow_html=True,
        )

        # Example download
        st.markdown(
            '<h3 style="color: #000000;">📥 Contoh Dokumen Panduan</h3>',
            unsafe_allow_html=True,
        )
        example_content = """PANDUAN FORMAT DOKUMEN

1. MARGIN
   - Atas: 4 cm
   - Bawah: 3 cm
   - Kiri: 4 cm
   - Kanan: 3 cm

2. FONT
   - Jenis: Times New Roman
   - Ukuran: 12 pt untuk isi
   - Ukuran: 14 pt untuk judul bab

3. SPASI
   - Antar baris: 1,5 (satu setengah)
   - Antar paragraf: 6 pt setelah paragraf

4. STRUKTUR DOKUMEN
   - Halaman Judul
   - Daftar Isi
   - BAB I PENDAHULUAN
   - BAB II TINJAUAN PUSTAKA
   - BAB III METODOLOGI
   - BAB IV HASIL DAN PEMBAHASAN
   - BAB V KESIMPULAN
   - Daftar Pustaka"""

        st.download_button(
            label="📄 Download Contoh",
            data=example_content,
            file_name="contoh_panduan.txt",
            mime="text/plain",
        )

    # Debug info di sidebar (opsional)
    with st.sidebar:
        if st.checkbox("🔧 Debug Info", value=False):
            st.write("**Environment:**")
            st.write(f"- Python: {sys.version[:20]}...")
            st.write(f"- Streamlit: {st.__version__}")
            st.write(f"- Working Dir: {os.getcwd()}")

            # Check API key
            try:
                api_key = st.secrets.get("GOOGLE_API_KEY") or os.getenv(
                    "GOOGLE_API_KEY"
                )
                if api_key:
                    st.success("✅ API Key tersedia")
                else:
                    st.error("❌ API Key tidak ditemukan")
            except:
                st.warning("⚠️ Error checking API key")

            # Check cache dir
            if os.path.exists("cache"):
                cache_files = len(os.listdir("cache"))
                st.info(f"📋 Cache files: {cache_files}")
            else:
                st.info("📋 Cache directory: tidak ada")


if __name__ == "__main__":
    main()
